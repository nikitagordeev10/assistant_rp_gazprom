'''
this module was made by Etienne Desautels and found at
http://etienned.github.io/posts/extract-text-from-word-docx-simply/
'''

# try:
#     from xml.etree.cElementTree import XML
# except ImportError:
#     from xml.etree.ElementTree import XML
# import zipfile
# import sys
#
# WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
# PARA = WORD_NAMESPACE + 'p'
# TEXT = WORD_NAMESPACE + 't'

import docx
import io
import pandas as pd

def get_docx_text(path):
    content = '\n'.join([p.text for p in docx.Document('data.docx').paragraphs])
    df = pd.read_csv(io.StringIO(content))
    return df
#     """
#     Take the path of a docx file as argument, return the text in unicode.
#     """
#     document = zipfile.ZipFile(path)
#     xml_content = document.read('word/document.xml')
#     document.close()
#     tree = XML(xml_content)
#
#     paragraphs = []
#     for paragraph in tree.iter(PARA):
#         texts = [node.text
#                  for node in paragraph.iter(TEXT)
#                  if node.text]
#         if texts:
#             paragraphs.append(''.join(texts))
#
#     return '\n\n'.join(paragraphs).encode('utf-8')

# if __name__ == '__main__':
#     # if len(sys.argv) != 2:
#     #     print ("Usage: python ", sys.argv[0], " <path of docx file>")
#     #     exit()
#     # txt = get_docx_text(sys.argv[1])
#
#     # content = docx.Document('data.docx').paragraphs[0].text
#     # or if all paragraphs
#     content = '\n'.join([p.text for p in docx.Document('data.docx').paragraphs])
#
#     df = pd.read_csv(io.StringIO(content))
#     print(df)